"""Build an editable Word (.docx) version of writeup.md, matching the
styling of the PDF (build_pdf.py) as closely as python-docx allows.
"""
import re
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "writeup.md"
OUT = "writeup.docx"

ACCENT = RGBColor(0x00, 0x72, 0xB2)
DARK = RGBColor(0x0B, 0x3D, 0x4D)
GRAY = RGBColor(0x55, 0x55, 0x55)
HEBREW_FONT = "Segoe UI"


def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    run.font.name = HEBREW_FONT
    r_fonts = rPr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        rPr.append(r_fonts)
    r_fonts.set(qn("w:cs"), HEBREW_FONT)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None, code=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if code:
        run.font.name = "Consolas"
    set_run_rtl(run)
    return run


def add_bottom_border(paragraph, color="0072B2", size=16):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def render_inline(paragraph, node, base_bold=False):
    """Render inline children (text, strong, em, code) of a soup node into a paragraph."""
    for child in node.children:
        if isinstance(child, str):
            if child.strip():
                add_run(paragraph, str(child), bold=base_bold)
        elif child.name == "strong":
            add_run(paragraph, child.get_text(), bold=True, color=DARK)
        elif child.name == "em":
            add_run(paragraph, child.get_text(), italic=True, color=GRAY)
        elif child.name == "code":
            add_run(paragraph, child.get_text(), code=True)
        else:
            add_run(paragraph, child.get_text(), bold=base_bold)


def add_field_label_or_paragraph(doc, p_tag):
    first = next((c for c in p_tag.children if not (isinstance(c, str) and not c.strip())), None)
    para = doc.add_paragraph()
    set_rtl(para)
    para.paragraph_format.space_after = Pt(6)
    if first is not None and getattr(first, "name", None) == "strong":
        label_para = doc.add_paragraph()
        set_rtl(label_para)
        label_para.paragraph_format.space_after = Pt(1)
        add_run(label_para, first.get_text(), bold=True, color=ACCENT, size=11)
        rest_para = doc.add_paragraph()
        set_rtl(rest_para)
        rest_para.paragraph_format.space_after = Pt(6)
        rest_children = list(p_tag.children)[list(p_tag.children).index(first) + 1:]
        for child in rest_children:
            if isinstance(child, str):
                if child.strip():
                    add_run(rest_para, str(child))
            elif child.name == "code":
                add_run(rest_para, child.get_text(), code=True)
            elif child.name == "em":
                add_run(rest_para, child.get_text(), italic=True, color=GRAY)
            elif child.name == "strong":
                add_run(rest_para, child.get_text(), bold=True, color=DARK)
            else:
                add_run(rest_para, child.get_text())
        doc._body._body.remove(para._p)
        return
    render_inline(para, p_tag)


def add_heading(doc, text, level=2):
    para = doc.add_paragraph()
    set_rtl(para)
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(8)
    add_bottom_border(para)
    add_run(para, text, bold=True, color=DARK, size=15)


def add_list(doc, list_tag, ordered=False):
    for li in list_tag.find_all("li", recursive=False):
        para = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        set_rtl(para)
        para.paragraph_format.space_after = Pt(3)
        render_inline(para, li)


def add_table(doc, table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    n_cols = len(rows[0].find_all(["th", "td"]))
    tbl = doc.add_table(rows=0, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    for r_idx, tr in enumerate(rows):
        row_cells = tbl.add_row().cells
        cells = tr.find_all(["th", "td"])
        is_header = tr.find("th") is not None
        for c_idx, cell in enumerate(cells):
            if c_idx >= n_cols:
                continue
            docx_cell = row_cells[c_idx]
            docx_cell.text = ""
            para = docx_cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(para)
            add_run(
                para,
                cell.get_text(),
                bold=is_header,
                color=RGBColor(0xFF, 0xFF, 0xFF) if is_header else None,
                size=10,
            )
            if is_header:
                shade_cell(docx_cell, "0072B2")
            elif r_idx % 2 == 0:
                shade_cell(docx_cell, "F4F8FA")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_image(doc, img_tag):
    src = img_tag.get("src", "")
    alt = img_tag.get("alt", "")
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(src, width=Cm(15))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(cap)
        add_run(cap, alt, italic=True, color=GRAY, size=9.5)
        cap.paragraph_format.space_after = Pt(10)
    except Exception as e:
        print(f"WARNING: could not embed image {src}: {e}")


def build():
    with open(SRC, encoding="utf-8") as f:
        md_text = f.read()

    html_body = markdown.markdown(md_text, extensions=["tables", "fenced_code", "nl2br"])
    soup = BeautifulSoup(html_body, "html.parser")

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.2)
    section.top_margin = section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = HEBREW_FONT
    style.font.size = Pt(11)

    first_h2_seen = False
    for node in soup.find_all(recursive=False):
        if node.name == "h1":
            para = doc.add_paragraph()
            set_rtl(para)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(6)
            add_run(para, node.get_text(), bold=True, color=DARK, size=22)
        elif node.name == "p" and node.find("em") and not first_h2_seen:
            para = doc.add_paragraph()
            set_rtl(para)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(14)
            add_run(para, node.get_text(), italic=True, color=GRAY, size=10.5)
        elif node.name == "h2":
            if node.get_text().strip() == "צוות":
                # Team section stays with the cover; render as a small heading.
                para = doc.add_paragraph()
                set_rtl(para)
                para.paragraph_format.space_before = Pt(4)
                add_run(para, "צוות הפרויקט", bold=True, color=ACCENT, size=12)
                first_h2_seen = False
                continue
            first_h2_seen = True
            add_heading(doc, node.get_text())
        elif node.name == "p":
            if node.find("img"):
                for img in node.find_all("img"):
                    add_image(doc, img)
            else:
                add_field_label_or_paragraph(doc, node)
        elif node.name == "ul":
            add_list(doc, node, ordered=False)
        elif node.name == "ol":
            add_list(doc, node, ordered=True)
        elif node.name == "table":
            add_table(doc, node)
        elif node.name == "div" and node.find("img"):
            for img in node.find_all("img"):
                add_image(doc, img)
        elif node.name == "hr":
            doc.add_page_break()

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
